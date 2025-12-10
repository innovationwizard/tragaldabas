"""Word document parser"""

from pathlib import Path
from typing import List

from core.models import (
    ReceptionResult, FileMetadata, SheetPreview
)
from core.enums import FileType
from core.exceptions import FileParseError
from .base import FileParser


class WordParser(FileParser):
    """Parser for Word documents (.docx)"""
    
    @property
    def supported_extensions(self) -> List[str]:
        return [".docx"]
    
    def detect_encoding(self, file_path: str) -> str:
        """Word files are binary, no encoding needed"""
        return "binary"
    
    def parse(self, file_path: str) -> ReceptionResult:
        """Parse Word document"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileParseError(f"File not found: {file_path}", file_path)
        
        try:
            from docx import Document
            
            doc = Document(path)
            file_size = path.stat().st_size
            
            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)
            
            # Split into lines for preview
            lines = full_text.split('\n')[:50]
            preview_rows = [[line] for line in lines]
            
            metadata = FileMetadata(
                file_path=str(path.absolute()),
                file_name=path.name,
                file_type=FileType.WORD_DOCX,
                file_size_bytes=file_size,
                encoding=None,
                page_count=len(doc.paragraphs)  # Approximate
            )
            
            preview = SheetPreview(
                sheet_name="Document",
                row_count=len(paragraphs),
                col_count=1,
                preview_rows=preview_rows,
                column_letters=["A"]
            )
            
            return ReceptionResult(
                metadata=metadata,
                previews=[preview],
                raw_data={"Document": full_text}  # Store as text
            )
            
        except ImportError:
            raise FileParseError(
                "python-docx not installed. Install with: pip install python-docx",
                file_path
            )
        except Exception as e:
            raise FileParseError(
                f"Failed to parse Word document: {e}",
                file_path
            ) from e

